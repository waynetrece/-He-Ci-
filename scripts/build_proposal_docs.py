#!/usr/bin/env python3
"""Build proposal-stage question Word files for HJ window and 凌越 window.

Generates:
  exports/HJ-提案階段確認問題-10主題.docx
  exports/凌越ERP-API對接-提案階段4題.docx

Source of truth: discussions/proposal-stage-questions-final.md (§A and §B)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "exports"
EXPORTS.mkdir(exist_ok=True)


def add_para(doc, text="", *, bold=False, size=None, after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_blank(doc):
    doc.add_paragraph()


def add_options(doc, options):
    for opt in options:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.add_run(opt)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.add_run(f"• {text}")


def add_reply_block(doc):
    p = doc.add_paragraph()
    p.add_run("回覆：").bold = True
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)


def build_hj():
    doc = Document()

    add_para(doc, "HJ 網站 — 提案階段確認問題清單（10 主題）",
             bold=True, size=18, after=4)
    add_para(doc, "提供給 HJ 窗口快速確認", size=11, after=2)
    add_para(doc, "日期：2026-05-02", size=11, after=8)

    add_para(doc, "說明", bold=True, size=13, after=4)
    add_para(doc,
             "本檔為提案階段的「最小確認題目集」，目的是把會影響網站系統架構、"
             "報價範圍與專案時程的關鍵業務規則先抓出來。每題建議由 HJ 內部相關窗口"
             "（業務 / 會計 / 倉管）討論後回覆即可。")
    add_para(doc,
             "未列入本檔的細節（如材積加總公式、各狀態起算日、發票欄位填寫規則等）"
             "會在簽約後 kickoff 階段再補完，這個階段不需要回覆。")
    add_para(doc, "如果某題目前還不確定，可以先標註「需內部確認」或留空。", after=8)

    questions = [
        (
            "一、樣品流程",
            "客戶同時加入樣品與公版商品時，HJ 希望採用哪一種：",
            ["A. 禁止一起結帳，需分次完成",
             "B. 同一頁送出，但系統自動拆成「樣品申請」與「一般訂單」兩筆",
             "C. 可一起付款，但後台拆單分開出貨"],
            None,
        ),
        (
            "二、私版詢價與 LINE",
            "私版詢價單成立後，是先只保留網站紀錄，等客服在 LINE 與客戶確認金額後，"
            "才同步進凌越報價／訂單嗎？",
            None,
            None,
        ),
        (
            "三、會員價格類別",
            "HJ 會員價格類別總共有幾種？是否只在固定的幾種價別內（例如直客 / 盤商 / 一般），"
            "還是會有「單一客戶專屬價」？",
            None,
            None,
        ),
        (
            "四、訂單同步凌越 / 新會員首單",
            "訂單送進凌越失敗時，是希望先成立網站訂單、由系統背景自動重送，"
            "還是直接擋下訂單請客戶稍後再試？",
            None,
            "另外，新會員第一次下單時若凌越裡尚無這位客戶的資料，HJ 希望網站先自動建立"
            "凌越客戶後再送訂單，還是先由網站成立訂單，待管理員審核 / 綁定客戶編號後再同步？",
        ),
        (
            "五、庫存同步速度",
            "HJ 對網站庫存更新的期待是哪一種：",
            ["A. 秒級即時（凌越改完，網站幾秒內就更新）",
             "B. 幾分鐘內，例如 5-10 分鐘",
             "C. 平時可有延遲，但客戶結帳前再次確認準確即可"],
            "不同方案的開發成本與技術風險差異較大，建議選擇前先了解三種對應的實作方式。",
        ),
        (
            "六、庫存顯示與預留庫存",
            "網站前台要顯示「實際庫存數量」、「有貨／缺貨」，還是「庫存不足 N 件以下時提示」？",
            None,
            "預留給連鎖店的庫存，可不可以在凌越用「獨立倉庫」管理？這樣網站只查一般倉的可用量，"
            "不需要再寫一套預留邏輯。",
        ),
        (
            "七、物流 / 材積 / 超商寄件單",
            "四大超商寄件單是否要在網站後台做批次列印？目前 HJ 是用哪個平台列印"
            "（綠界物流、超商自有平台、人工）？",
            None,
            "HJ 是否已經有一套判斷超商 / 宅配可配送的材積與重量規則？若還沒有，"
            "是否接受由網站採保守規則：超過超商限制就只顯示宅配，超過宅配限制就提示聯絡客服？",
        ),
        (
            "八、金流與發票",
            "綠界帳號由誰申請？何時可提供測試與正式商店代號？",
            None,
            "發票由哪一方開立：凌越、綠界、另接電子發票服務，還是會計人工開立？",
        ),
        (
            "九、退貨退款",
            "退貨退款時，目前的處理方式是哪一種：",
            ["A. 全部由會計人工處理（信用卡退刷、匯款退款、發票作廢分開做）",
             "B. 希望系統連動（網站申請 → 自動退刷 → 自動作廢發票 → 自動更新凌越）",
             "C. 混合（部分連動、部分人工）"],
            None,
        ),
        (
            "十、商品資料來源與訂單狀態",
            "商品資料的正式來源是哪一個：凌越、Shopline Excel、新網站後台，還是混合維護？",
            None,
            "訂單狀態（待付款／已付款／備貨中／已出貨／已完成／已取消）是由 HJ 人員在網站後台"
            "更新，還是希望凌越狀態異動後自動反映到網站？",
        ),
    ]

    for title, q_main, options, q_followup in questions:
        add_para(doc, title, bold=True, size=14, after=4)
        add_para(doc, q_main)
        if options:
            add_options(doc, options)
        if q_followup:
            add_blank(doc)
            add_para(doc, q_followup)
        add_reply_block(doc)

    add_para(doc, "備註", bold=True, size=12, after=4)
    add_para(doc,
             "本檔總共 10 題，建議由 HJ 內部各窗口分工填寫，"
             "彙整後一次回覆即可。回覆完後可以直接在 Word 上編輯儲存，"
             "或印出後手寫掃描傳回。")

    out = EXPORTS / "HJ-提案階段確認問題-10主題.docx"
    doc.save(str(out))
    print(f"Wrote {out}")


def build_lingyue():
    doc = Document()

    add_para(doc, "凌越 ERP API 對接 — 提案階段確認問題（4 題）",
             bold=True, size=18, after=4)
    add_para(doc, "對象：凌越資訊技術窗口", size=11, after=2)
    add_para(doc, "日期：2026-05-02", size=11, after=8)

    add_para(doc, "說明", bold=True, size=13, after=4)
    add_para(doc,
             "馬亞團隊（HJ 新網站開發方）已收到 HJ 提供的兩份凌越 ERP API 文件"
             "（5 頁元件流程說明與 207 頁資料清單）。")
    add_para(doc,
             "本檔僅為提案階段的最小確認題集，用以判斷 API 文件是否能撐起 HJ 專案需求，"
             "以及合作確認後需多久才能進入實作。")
    add_para(doc,
             "完整的欄位對照、XML 範例、單號規則、錯誤碼等技術細節，"
             "會在合作確認後另行請貴司提供，這個階段不需要回覆。", after=8)

    questions = [
        (
            "一、授權範圍",
            "HJ 帳號目前授權哪些 API 資料類別？是否包含「商品、客戶、訂單、庫存、報價、"
            "收款、發票」這七類？若有未授權，是否可在合作後開通？",
            None,
        ),
        (
            "二、客戶別商品價格表",
            "凌越是否有「客戶別商品價格表」或「客戶貨品項號對照」這類 API？"
            "若有，請簡述欄位或提供樣例。",
            "此題影響 HJ 「會員分級價」是否能由 ERP 直接承接，是 HJ 專案的重要前置確認點。",
        ),
        (
            "三、啟動時程",
            "合作確認之後，凌越多久可以提供以下對接資料：",
            "此題用來估算 HJ 專案 kickoff 後的前置時間。",
        ),
        (
            "四、整合窗口與經驗",
            "凌越是否有固定的技術對接窗口或既有的第三方系統串接流程？"
            "HJ 過去是否曾透過貴司 API 與其他第三方系統（電商、CRM 等）整合過？",
            "此題用來評估溝通成本與整合風險，不需提供 SLA 或正式合約細節。",
        ),
    ]

    for i, (title, q_main, note) in enumerate(questions):
        add_para(doc, title, bold=True, size=14, after=4)
        add_para(doc, q_main)
        if i == 2:
            # 啟動時程 needs a bullet list
            add_bullet(doc, "測試環境（endpoint / WSDL）")
            add_bullet(doc, "測試帳號、密碼、公司代號")
            add_bullet(doc, "商品 / 客戶 / 訂單 / 庫存等核心流程的可執行 XML 範例")
        if note:
            add_blank(doc)
            add_para(doc, note)
        add_reply_block(doc)

    out = EXPORTS / "凌越ERP-API對接-提案階段4題.docx"
    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    build_hj()
    build_lingyue()
